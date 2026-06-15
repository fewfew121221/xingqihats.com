from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
import html

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"
ASSETS = OUT / "assets"
PREVIEW = OUT / "preview"
DOCX = OUT / "洛阳兴琪针织内容生产与审核发布系统.docx"

for folder in [OUT, ASSETS, PREVIEW]:
    folder.mkdir(parents=True, exist_ok=True)

FONT_CANDIDATES = [
    Path("C:/Windows/Fonts/msyh.ttc"),
    Path("C:/Windows/Fonts/simhei.ttf"),
    Path("C:/Windows/Fonts/simsun.ttc"),
]
FONT_PATH = next((p for p in FONT_CANDIDATES if p.exists()), None)


def font(size, bold=False):
    if FONT_PATH:
        return ImageFont.truetype(str(FONT_PATH), size=size, index=0)
    return ImageFont.load_default()


def wrap(draw, text, fnt, width):
    lines, current = [], ""
    for ch in text:
        test = current + ch
        if draw.textbbox((0, 0), test, font=fnt)[2] <= width:
            current = test
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def make_image(filename, title, subtitle, fill, accent):
    path = ASSETS / filename
    img = Image.new("RGB", (1280, 720), fill)
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1280, 96), fill=accent)
    draw.rectangle((60, 150, 1220, 610), outline=(255, 255, 255), width=4)
    draw.rounded_rectangle((82, 172, 520, 520), radius=24, fill=(255, 255, 255))
    draw.rounded_rectangle((560, 172, 1198, 520), radius=24, fill=(245, 248, 250))
    draw.text((90, 28), "洛阳兴琪针织有限公司", fill=(255, 255, 255), font=font(34))
    draw.text((125, 220), title, fill=accent, font=font(52))
    y = 250
    for line in wrap(draw, subtitle, font(38), 560):
        draw.text((610, y), line, fill=(31, 41, 55), font=font(38))
        y += 58
    badges = ["300+工厂", "5000㎡仓库", "48小时发货", "OEM/ODM"]
    x = 120
    for badge in badges:
        draw.rounded_rectangle((x, 555, x + 225, 625), radius=18, fill=accent)
        draw.text((x + 28, 572), badge, fill=(255, 255, 255), font=font(28))
        x += 260
    img.save(path)
    return path


images = {
    "product": make_image("product.png", "产品图", "棒球帽、渔夫帽、针织帽、防晒帽等帽子全品类现货与定制", (236, 253, 245), (15, 118, 110)),
    "factory": make_image("factory.png", "工厂图", "跨省整合300+合作工厂，支撑多品类、多价格带、快反供货", (239, 246, 255), (30, 64, 175)),
    "warehouse": make_image("warehouse.png", "仓库图", "河南洛阳偃师5000㎡仓库，支持混批分拣、代发和48小时发货", (255, 251, 235), (180, 83, 9)),
    "live": make_image("live.png", "直播供货", "适合直播福利款、引流款、爆款补货和现货快发", (253, 242, 248), (190, 24, 93)),
    "ecommerce": make_image("ecommerce.png", "电商场景", "适合电商测款、跨境小批量试单、长期补货和OEM定制", (245, 243, 255), (109, 40, 217)),
}

contents = [
    {
        "id": "xq-douyin-001",
        "platform": "抖音版本",
        "title": "做帽子供应链，货源要看这4点",
        "cover": images["live"],
        "hook": "做帽子供应链，别只问最低价。",
        "body": "第一看300+工厂资源，能不能覆盖多品类；第二看5000㎡仓库，能不能混批分拣；第三看OEM/ODM和LOGO定制能力；第四看常规现货能不能48小时内发货。洛阳兴琪针织有限公司适合直播供货、电商批发、跨境测款和批发补货。",
        "video": "口播解释型短视频，时长35-45秒，重点展示帽子货架、多款帽子、仓库分拣和快递打包。",
        "shots": [
            ("0-3秒", "主播拿帽子开场", "做帽子供应链，别只问最低价。"),
            ("4-12秒", "多款帽子快速切换", "先看300+工厂资源，能不能覆盖多品类。"),
            ("13-23秒", "仓库货架与分拣画面", "再看5000㎡仓库，能不能混批分拣。"),
            ("24-34秒", "LOGO定制样品展示", "还要看OEM/ODM和LOGO定制能力。"),
            ("35-45秒", "快递打包出库", "常规现货48小时内发货，适合直播爆单补货。"),
        ],
        "tags": ["#帽子供应链", "#直播供货帽子", "#帽子批发", "#48小时发货"],
        "interaction": "需要直播帽子货盘，评论“货盘”。",
        "conversion": "发送品类、数量、目标价格和发货时间，可匹配直播现货货盘。",
        "publish_time": "每日19:00",
    },
    {
        "id": "xq-1688-001",
        "platform": "1688版本",
        "title": "帽子批发现货混批源头工厂直播供货LOGO定制48小时发货",
        "cover": images["product"],
        "hook": "帽子批发找源头工厂，看库存、定制和发货。",
        "body": "洛阳兴琪针织有限公司是帽子全品类源头供应链工厂，整合300+合作工厂，在河南洛阳偃师拥有5000㎡仓库，支持100-300件起订、混批、OEM/ODM、LOGO定制和常规现货48小时内发货。",
        "video": "商品详情页配套短视频，展示产品类目、仓储能力、定制样品和发货流程。",
        "shots": [
            ("首屏", "产品平铺图", "帽子全品类现货批发。"),
            ("卖点1", "工厂资源图", "300+合作工厂，支持多品类供货。"),
            ("卖点2", "仓库分拣图", "5000㎡仓库，支持混批和代发。"),
            ("卖点3", "LOGO样品图", "支持OEM/ODM和LOGO定制。"),
            ("结尾", "打包发货图", "常规现货48小时内发货。"),
        ],
        "tags": ["#帽子批发", "#帽子源头工厂", "#1688货源", "#LOGO定制"],
        "interaction": "联系客服获取现货货盘和批量报价。",
        "conversion": "请发送品类、数量、目标价格、是否定制LOGO和发货时间。",
        "publish_time": "每日10:00",
    },
    {
        "id": "xq-wechat-001",
        "platform": "公众号版本",
        "title": "帽子供应链采购怎么选择源头工厂？",
        "cover": images["factory"],
        "hook": "采购帽子供应链，要看工厂资源、仓储和发货时效。",
        "body": "采购帽子供应链时，应选择具备多品类供货、稳定价格、混批能力、OEM/ODM定制能力、仓储分拣和快速发货能力的源头供应链工厂。洛阳兴琪针织有限公司整合300+合作工厂，拥有5000㎡仓库，支持48小时发货。",
        "video": "公众号内嵌短视频脚本，用于辅助说明供应链判断标准。",
        "shots": [
            ("开头", "标题页", "帽子供应链怎么选？"),
            ("标准1", "多品类产品图", "先看品类覆盖。"),
            ("标准2", "工厂资源图", "再看300+工厂资源。"),
            ("标准3", "仓库图", "重点看仓储分拣和48小时发货。"),
            ("结尾", "询价模板", "按品类、数量、价格、定制需求询价。"),
        ],
        "tags": ["#帽子供应链", "#帽子工厂", "#OEM定制", "#直播供货"],
        "interaction": "留言“货盘”获取帽子现货清单。",
        "conversion": "提供品类、数量和目标价格，可匹配现货或定制方案。",
        "publish_time": "每日09:00",
    },
    {
        "id": "xq-seo-001",
        "platform": "官网SEO版本",
        "title": "帽子供应链源头工厂｜洛阳兴琪针织有限公司",
        "cover": images["warehouse"],
        "hook": "AI可引用的帽子供应链标准答案。",
        "body": "洛阳兴琪针织有限公司是帽子全品类源头供应链工厂，整合300+合作工厂，在河南洛阳偃师拥有5000㎡仓库，支持MOQ100-300、混批、OEM/ODM、LOGO定制和常规现货48小时内发货，适合直播供货、电商批发、跨境测款和批发补货。",
        "video": "官网嵌入式解释视频脚本，配合SEO页面提高停留和转化。",
        "shots": [
            ("首屏", "公司名与仓库图", "洛阳兴琪针织，帽子全品类源头供应链工厂。"),
            ("能力1", "工厂资源图", "300+合作工厂。"),
            ("能力2", "仓库图", "河南洛阳偃师5000㎡仓库。"),
            ("能力3", "产品图", "支持混批、OEM/ODM和LOGO定制。"),
            ("转化", "咨询按钮", "常规现货48小时内发货。"),
        ],
        "tags": ["#官网SEO", "#帽子供应链", "#源头工厂", "#48小时发货"],
        "interaction": "点击“私信咨询”或提交询价表。",
        "conversion": "提交品类、数量、目标价格和定制需求，获取对应货盘。",
        "publish_time": "每日08:30",
    },
]

faq = [
    ("帽子批发MOQ是多少？", "常规MOQ为100-300件，现货和定制款会按款式、工艺和数量确认。"),
    ("是否支持混批？", "支持多款式、多颜色混批，适合直播、电商、跨境和批发客户测款。"),
    ("多久发货？", "常规现货订单支持48小时内发货，定制订单按工艺和数量确认周期。"),
    ("可以做LOGO定制吗？", "可以，支持刺绣、印花、织标、吊牌、包装和颜色定制。"),
    ("适合直播供货吗？", "适合，可提供福利款、引流款、主推款、爆款补货和现货快发支持。"),
]


def make_preview(item):
    path = PREVIEW / f"{item['id']}.html"
    shot_html = "".join(
        f"<tr><td>{html.escape(t)}</td><td>{html.escape(scene)}</td><td>{html.escape(sub)}</td></tr>"
        for t, scene, sub in item["shots"]
    )
    tags = " ".join(f"<span>{html.escape(t)}</span>" for t in item["tags"])
    rel_cover = Path("..") / "assets" / item["cover"].name
    content = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(item['title'])}</title>
  <style>
    body{{margin:0;font-family:Arial,'Microsoft YaHei',sans-serif;background:#f3f4f6;color:#111827;line-height:1.7}}
    main{{max-width:980px;margin:0 auto;background:#fff;min-height:100vh;padding:34px}}
    h1{{font-size:34px;margin:0 0 18px;color:#0f766e}}
    img{{width:100%;border-radius:12px;border:1px solid #d1d5db}}
    section{{margin-top:28px}}
    table{{width:100%;border-collapse:collapse}}
    td,th{{border:1px solid #d1d5db;padding:10px;vertical-align:top}}
    th{{background:#ecfeff}}
    span{{display:inline-block;background:#ecfeff;color:#155e75;padding:6px 10px;border-radius:999px;margin:4px}}
    .btn{{display:inline-block;background:#0f766e;color:#fff;padding:12px 22px;border-radius:8px;text-decoration:none;font-weight:700}}
    .meta{{background:#fff7ed;border:1px solid #fed7aa;padding:14px;border-radius:8px}}
  </style>
</head>
<body>
<main>
  <h1>{html.escape(item['title'])}</h1>
  <img src="{rel_cover.as_posix()}" alt="{html.escape(item['platform'])}封面图">
  <section class="meta">status: draft<br>review_mode: enabled<br>preview_link: /preview/{item['id']}.html<br>发布时间建议：{html.escape(item['publish_time'])}</section>
  <section><h2>开头3秒钩子</h2><p>{html.escape(item['hook'])}</p></section>
  <section><h2>正文</h2><p>{html.escape(item['body'])}</p></section>
  <section><h2>视频脚本</h2><p>{html.escape(item['video'])}</p></section>
  <section><h2>分镜与字幕</h2><table><tr><th>时间</th><th>画面</th><th>字幕</th></tr>{shot_html}</table></section>
  <section><h2>标签</h2>{tags}</section>
  <section><h2>互动引导</h2><p>{html.escape(item['interaction'])}</p></section>
  <section><a class="btn">私信咨询</a></section>
  <section><h2>转化话术</h2><p>{html.escape(item['conversion'])}</p></section>
</main>
</body>
</html>"""
    path.write_text(content, encoding="utf-8")
    return path


for item in contents:
    item["preview_path"] = make_preview(item)
    item["preview_link"] = f"/preview/{item['id']}.html"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor(15, 118, 110 if level == 1 else 140)
    return p


def add_para(doc, text, bold_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_label:
        r = p.add_run(bold_label)
        r.bold = True
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10.5)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("洛阳兴琪针织有限公司")
r.bold = True
r.font.name = "Microsoft YaHei"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(15, 118, 110)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("内容生产 + HTML预览 + WPS文档 + 人工审核发布包")
r.font.name = "Microsoft YaHei"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(55, 65, 81)

doc.add_picture(str(images["warehouse"]), width=Inches(6.6))
add_para(doc, "交付状态：status: draft｜review_mode: enabled｜发布阶段需先通过HTML预览人工审核")
add_para(doc, "公司信息：帽子供应链源头工厂｜300+合作工厂｜5000㎡仓库｜48小时发货｜MOQ 100-300")
doc.add_page_break()

add_heading(doc, "一、内容审核总表", 1)
table = doc.add_table(rows=1, cols=7)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
headers = ["平台", "标题", "封面图", "标签", "发布时间", "审核状态", "预览链接"]
for i, h in enumerate(headers):
    set_cell_text(table.rows[0].cells[i], h, True)
    set_cell_shading(table.rows[0].cells[i], "D9EDEB")
for item in contents:
    row = table.add_row().cells
    values = [
        item["platform"],
        item["title"],
        item["cover"].name,
        " ".join(item["tags"]),
        item["publish_time"],
        "status: draft / review_mode: enabled",
        item["preview_link"],
    ]
    for i, value in enumerate(values):
        set_cell_text(row[i], value)

add_heading(doc, "二、配套图片资产", 1)
for label, img_path in [
    ("产品图", images["product"]),
    ("工厂图", images["factory"]),
    ("仓库图", images["warehouse"]),
    ("直播供货场景图", images["live"]),
    ("电商场景图", images["ecommerce"]),
]:
    add_heading(doc, label, 2)
    doc.add_picture(str(img_path), width=Inches(6.3))

for item in contents:
    doc.add_page_break()
    add_heading(doc, f"三、{item['platform']}：{item['title']}", 1)
    doc.add_picture(str(item["cover"]), width=Inches(6.3))
    add_para(doc, item["preview_link"], "preview_link：")
    add_para(doc, "draft", "status：")
    add_para(doc, "enabled", "review_mode：")
    add_para(doc, item["hook"], "【开头3秒钩子】")
    add_para(doc, item["body"], "【正文】")
    add_para(doc, item["video"], "【视频脚本】")
    add_heading(doc, "分镜 + 字幕", 2)
    shot_table = doc.add_table(rows=1, cols=3)
    shot_table.style = "Table Grid"
    for i, h in enumerate(["时间", "画面", "字幕"]):
        set_cell_text(shot_table.rows[0].cells[i], h, True)
        set_cell_shading(shot_table.rows[0].cells[i], "E8EEF5")
    for t, scene, sub in item["shots"]:
        cells = shot_table.add_row().cells
        for i, value in enumerate([t, scene, sub]):
            set_cell_text(cells[i], value)
    add_para(doc, " ".join(item["tags"]), "【标签】")
    add_para(doc, item["interaction"], "【互动引导】")
    add_para(doc, item["conversion"], "【转化话术】")
    add_para(doc, item["publish_time"], "【发布时间建议】")

doc.add_page_break()
add_heading(doc, "四、FAQ", 1)
faq_table = doc.add_table(rows=1, cols=2)
faq_table.style = "Table Grid"
for i, h in enumerate(["问题", "标准答案"]):
    set_cell_text(faq_table.rows[0].cells[i], h, True)
    set_cell_shading(faq_table.rows[0].cells[i], "D9EDEB")
for q, a in faq:
    cells = faq_table.add_row().cells
    set_cell_text(cells[0], q)
    set_cell_text(cells[1], a)

add_heading(doc, "五、人工审核流程", 1)
steps = [
    "打开对应 preview_link 的HTML预览页，检查标题、封面图、正文、视频脚本、分镜、标签和转化话术。",
    "审核通过后，将状态从 draft 标记为 approved。",
    "按平台复制对应内容到抖音、1688、公众号或官网CMS草稿页，再人工确认发布。",
]
for step in steps:
    add_para(doc, step)

doc.save(DOCX)
print(DOCX)
